import os
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from litellm import completion
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# SYNCED DATA MODEL (Matches all 19 UI fields)
class ProjectDetails(BaseModel):
    project_name: str = "N/A"
    geo_scope: str = "N/A"
    lawful_basis: str = "N/A"
    data_subjects: str = "N/A"
    collection_source: str = "N/A"
    project_desc: str = "N/A"
    data_collected: str = "N/A"
    transparency_measures: str = "N/A"
    data_minimization: str = "N/A"
    data_quality: str = "N/A"
    security_measures: str = "N/A"
    third_parties: str = "N/A"
    intl_transfers: str = "N/A"
    retention_period: str = "N/A"
    retention_policy: str = "N/A"
    individual_rights: str = "N/A"
    initial_risk: str = "Medium"
    license_key: str = "FREE"

class FinalReportRequest(ProjectDetails):
    identified_risks: str = "N/A"

def add_formatted_text_to_word(doc, text):
    lines = text.split('\n')
    in_table = False
    table = None
    for line in lines:
        line = line.strip()
        if not line:
            in_table = False
            continue
        if line.startswith('|'):
            parts = line.split('|')
            cells = [p.strip() for p in parts[1:-1]]
            if all(all(c in '-: ' for c in cell) for cell in cells) and len(cells) > 0:
                continue
            if not in_table:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    hdr_cells[i].text = cell_text.replace('**', '').strip()
                    for paragraph in hdr_cells[i].paragraphs:
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.bold = True
                in_table = True
            else:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text.replace('**', '').replace('<br>', '\n').strip()
            continue
        else:
            in_table = False
        if line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True

@app.get("/")
def home():
    return {"message": "Global DPIA Engine Online"}

@app.post("/api/analyze")
async def analyze_risks(data: ProjectDetails):
    prompt = f"""
    Act as a Global Privacy Counsel. Conduct a high-level risk assessment for:
    Project: {data.project_name} | Scope: {data.geo_scope}
    Identify 3-5 key privacy risks and mitigation requirements.
    """
    try:
        response = completion(model="gemini/gemini-2.5-flash", messages=[{"role": "user", "content": prompt}])
        return {"status": "success", "risks": response.choices[0].message.content}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/api/generate-report")
async def generate_final_report(data: FinalReportRequest):
    try:
        doc = Document()
        title = doc.add_heading('DATA PROTECTION IMPACT ASSESSMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Project Context Table
        doc.add_heading('1. Information Asset & Processing Context', level=1)
        meta = doc.add_table(rows=7, cols=2)
        meta.style = 'Table Grid'
        items = [
            ("Project Name", data.project_name),
            ("Geographic Scope", data.geo_scope),
            ("Lawful Basis", data.lawful_basis),
            ("Data Subjects", data.data_subjects),
            ("Collection Source", data.collection_source),
            ("Retention Period", data.retention_period),
            ("Initial Risk Level", data.initial_risk)
        ]
        for i, (label, val) in enumerate(items):
            meta.cell(i,0).text = label
            meta.cell(i,0).paragraphs[0].runs[0].bold = True
            meta.cell(i,1).text = str(val)

        prompt = f"""
        Act as a Senior Global Privacy Consultant. Write a formal DPIA Report for '{data.project_name}'.
        Details: {data.project_desc} | Data: {data.data_collected} | Security: {data.security_measures} | Transfers: {data.intl_transfers}
        Rights: {data.individual_rights} | Retention: {data.retention_period}
        
        Structure the report exactly as follows:
        ## 2. Executive Summary
        (Summarize the activity and overall compliance posture)
        
        ## 3. Scope of Processing
        (Define the nature, scale, and context of the processing activity)
        
        ## 4. Data Governance Analysis Table
        You MUST provide the governance analysis in a Markdown table:
        | Governance Area | Finding | Recommendation | Comment |
        (For Comment, leave it empty. Areas: Transparency, Minimization, Quality)
        
        ## 5. Formal Risk Assessment Matrix
        You MUST provide the risk matrix in a Markdown table:
        | Privacy Risk | Recommended Mitigation | Audit Evidence Required | Comment |
        (For Comment, leave it empty)
        
        ## 6. Final Risk Rating & Justification
        Write this section exactly following this structure:
        - **Initial Risk Rating:** {data.initial_risk}
        - (Explain why the project was inherently risky before mitigations)
        - **Residual Risk Rating (Post-Mitigation):** [Assess as Acceptable/Low/Medium]
        - **Justification:** Provide a detailed numbered list covering:
            1. Legal Foundations
            2. Data Lifecycle Management
            3. Robust Security
            4. Transparency and Empowerment
            5. Cross-Border Compliance
            6. Special Safeguards (e.g., Employees or Children)
        - (Concluding statement on project proceeding)
        
        Tone: Corporate, Authoritative, Neutral. No mention of AI.
        """
        
        response = completion(model="gemini/gemini-2.5-flash", messages=[{"role": "user", "content": prompt}])
        add_formatted_text_to_word(doc, response.choices[0].message.content)
        
        doc.add_paragraph("\nReport generated by DPIA Enterprise Systems. Valid for global compliance verification.")
        file_path = f"/tmp/Report.docx"
        doc.save(file_path)
        return FileResponse(path=file_path, filename=f"Global_DPIA_{data.project_name}.docx")
    except Exception as e:
        return {"status": "error", "message": str(e)}
